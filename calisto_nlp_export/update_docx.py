from docx import Document

doc = Document('knowledge_base/faq_customer_support_calisto.docx')
doc.add_paragraph('Q: What is your refund or return policy?')
doc.add_paragraph('A: Calisto offers a 14-day return and refund policy for all standard eyewear items provided they are in their original unworn condition with the original receipt. If you are not satisfied with your purchase, bring the item back to any Calisto store location to process a prompt refund or exchange. For custom prescription lenses, refunds are assessed on a case-by-case basis.')
doc.save('knowledge_base/faq_customer_support_calisto.docx')
